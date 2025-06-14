import os
from docx import Document

def extract_hyperlinks(doc_path):
    links = []
    try:
        doc = Document(doc_path)
        for para in doc.paragraphs:
            for run in para.runs:
                if "HYPERLINK" in run._element.xml:
                    xml = run._element.xml
                    start = xml.find("HYPERLINK") + len("HYPERLINK ")
                    end = xml.find('"', start)
                    if start != -1 and end != -1:
                        link = xml[start:end].strip().replace('"', '')
                        links.append(link)

        # Check inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if "HYPERLINK" in run._element.xml:
                                xml = run._element.xml
                                start = xml.find("HYPERLINK") + len("HYPERLINK ")
                                end = xml.find('"', start)
                                if start != -1 and end != -1:
                                    link = xml[start:end].strip().replace('"', '')
                                    links.append(link)
    except Exception as e:
        print(f"❌ Error reading {doc_path}: {e}")
    return links

def find_docs_with_links(root_folder):
    for foldername, subfolders, filenames in os.walk(root_folder):
        for filename in filenames:
            if filename.lower().endswith(".docx"):
                full_path = os.path.join(foldername, filename)
                links = extract_hyperlinks(full_path)
                if links:
                    print(f"\n🔗 {full_path}")
                    for i, link in enumerate(set(links), start=1):
                        print(f"   {i}. {link}")

# 🔍 Replace this with your actual folder path
root_dir = r"C:\Users\hp\OneDrive\Desktop\epchatbot-finalvr-main - Copy\user manual"
find_docs_with_links(root_dir)
